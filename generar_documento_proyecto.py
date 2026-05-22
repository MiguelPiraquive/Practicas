# -*- coding: utf-8 -*-
"""
Documento de Práctica Profesional — SIGET-HC
Formato ICONTEC + APA, Universidad Piloto de Colombia (Alto Magdalena).
- Usa estilos Heading 1/2 (la tabla de contenido se genera automática en Word).
- Incluye citas APA en el cuerpo y referencias al final.
- Mantiene textualmente la problemática y los objetivos del estudiante.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Arial"


def set_default_font(doc):
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3",
                       "Title", "Subtitle"]:
        try:
            style = doc.styles[style_name]
            style.font.name = FONT
            rpr = style.element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:ascii"), FONT)
            rfonts.set(qn("w:hAnsi"), FONT)
            rfonts.set(qn("w:cs"), FONT)
        except KeyError:
            pass
    doc.styles["Normal"].font.size = Pt(12)


def set_margins(section):
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(2)


def par(doc, text="", *, bold=False, align=None, size=12, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    if text:
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
    return p


def justified(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(12)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(12)
    return p


def heading1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(18)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text.upper())
    r.font.name = FONT
    r.font.size = Pt(14)
    r.bold = True


def heading2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text.upper())
    r.font.name = FONT
    r.font.size = Pt(12)
    r.bold = True


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def insert_toc(doc):
    """Inserta campo TOC. Word lo poblará automáticamente al abrir
    el documento (clic derecho → Actualizar campos, o F9)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = ('Tabla de contenido — clic derecho sobre esta línea '
                        'en Word y elija "Actualizar campos" (F9).')
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld1)
    r.append(instr)
    r.append(fld2)
    r.append(placeholder)
    r.append(fld3)


# ═══════════════════════════════════════════════════════════════
doc = Document()
set_default_font(doc)
for s in doc.sections:
    set_margins(s)

# ── CUBIERTA ───────────────────────────────────────────────────
for _ in range(3):
    par(doc)
par(doc, "SISTEMA DE GESTIÓN Y TRAZABILIDAD DE HISTORIAS CLÍNICAS (SIGET-HC)",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
for _ in range(8):
    par(doc)
par(doc, "MIGUEL ÁNGEL PIRAQUIVE PACHÓN",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(10):
    par(doc)
par(doc, "UNIVERSIDAD PILOTO DE COLOMBIA",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
par(doc, "SECCIONAL DEL ALTO MAGDALENA",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
par(doc, "FACULTAD DE INGENIERÍA",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
par(doc, "PROGRAMA DE INGENIERÍA DE SISTEMAS",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
par(doc, "GIRARDOT, CUNDINAMARCA",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
par(doc, "2026", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
page_break(doc)

# ── PORTADA INTERIOR ───────────────────────────────────────────
for _ in range(3):
    par(doc)
par(doc, "SISTEMA DE GESTIÓN Y TRAZABILIDAD DE HISTORIAS CLÍNICAS (SIGET-HC)",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
for _ in range(4):
    par(doc)
par(doc, "MIGUEL ÁNGEL PIRAQUIVE PACHÓN",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
par(doc, "(Correo electrónico institucional)",
    align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(3):
    par(doc)
par(doc, "Proyecto presentado para el curso de Práctica Profesional.",
    align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(3):
    par(doc)
par(doc, "Docente", align=WD_ALIGN_PARAGRAPH.CENTER)
par(doc, "NOMBRE DEL DOCENTE A CARGO DEL CURSO",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
par(doc, "(Título profesional)", align=WD_ALIGN_PARAGRAPH.CENTER)
par(doc, "(Correo electrónico institucional)",
    align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(5):
    par(doc)
par(doc, "UNIVERSIDAD PILOTO DE COLOMBIA",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
par(doc, "SECCIONAL DEL ALTO MAGDALENA",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
par(doc, "FACULTAD DE INGENIERÍA",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
par(doc, "PROGRAMA DE INGENIERÍA DE SISTEMAS",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
par(doc, "GIRARDOT, CUNDINAMARCA",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
par(doc, "2026", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
page_break(doc)

# ── TABLA DE CONTENIDO ─────────────────────────────────────────
par(doc, "CONTENIDO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
par(doc)
insert_toc(doc)
par(doc)
par(doc,
    "Nota: ubíquese sobre la tabla anterior en Word, presione clic "
    "derecho y elija \"Actualizar campos\" (F9) para que se carguen "
    "los títulos y los números de página automáticamente.",
    italic=True, size=10)
page_break(doc)

# ── INTRODUCCIÓN ───────────────────────────────────────────────
heading1(doc, "Introducción")
justified(doc,
    "El presente documento expone el desarrollo del proyecto Sistema de "
    "Gestión y Trazabilidad de Historias Clínicas (SIGET-HC), elaborado "
    "en el marco de la práctica profesional adelantada en la Clínica "
    "Junical, ubicada en la ciudad de Girardot, Cundinamarca. El "
    "proyecto se concibe como un sistema de información web orientado "
    "a modernizar el proceso administrativo de solicitud, control y "
    "entrega de historias clínicas, sustituyendo el manejo manual "
    "mediante archivos de hoja de cálculo por una herramienta "
    "organizada, auditable y segura (Laudon y Laudon, 2020).")
justified(doc,
    "Los antecedentes que motivan este trabajo se sustentan en estudios "
    "previos que evidencian que la sistematización de procesos en "
    "instituciones prestadoras de servicios de salud reduce los tiempos "
    "de atención, mejora la integridad de los registros y facilita el "
    "cumplimiento normativo (Pressman y Maxim, 2020). En Colombia, el "
    "manejo de la historia clínica está regulado por la Resolución 1995 "
    "de 1999 del Ministerio de Salud, modificada posteriormente por la "
    "Resolución 839 de 2017, que establecen los principios de "
    "integralidad, secuencialidad, racionalidad científica, "
    "disponibilidad y oportunidad (Ministerio de Salud y Protección "
    "Social, 1999, 2017).")
justified(doc,
    "El objetivo central consiste en diseñar, desarrollar e implementar "
    "una solución informática que optimice el registro, control y "
    "seguimiento de las solicitudes de historias clínicas, garantizando "
    "la integridad de los datos y agilizando la atención en ventanilla. "
    "El alcance comprende el análisis del proceso actual, el diseño de "
    "la arquitectura, el desarrollo de los módulos principales y la "
    "implementación en el entorno real de la institución.")
justified(doc,
    "La metodología empleada corresponde a un enfoque de investigación "
    "aplicada con un ciclo de vida iterativo e incremental, apoyado en "
    "técnicas de levantamiento de requerimientos, diseño centrado en el "
    "usuario y pruebas funcionales (Sommerville, 2021). El significado "
    "del estudio radica en su aporte directo a la calidad de un servicio "
    "administrativo crítico, contribuyendo al Objetivo de Desarrollo "
    "Sostenible número tres relativo a la Salud y el Bienestar "
    "(Organización de las Naciones Unidas, 2015).")
page_break(doc)

# ── 1. DEFINICIÓN DEL PROBLEMA ─────────────────────────────────
heading1(doc, "1. Definición del problema")
justified(doc,
    "Actualmente, el control de entregas se realiza mediante archivos "
    "planos (Excel) sin validación de datos, lo que genera riesgos de "
    "inconsistencia en la información, duplicidad de registros, "
    "dificultad para generar reportes históricos y falta de seguridad "
    "sobre quién modifica los datos.")
justified(doc,
    "Esta situación coincide con lo reportado por diversos autores, "
    "quienes señalan que la dependencia de herramientas ofimáticas "
    "para la gestión documental en entornos hospitalarios produce "
    "errores humanos, ausencia de trazabilidad y pérdida de control "
    "sobre la información (Laudon y Laudon, 2020). Adicionalmente, el "
    "manejo manual compromete los principios de integridad y "
    "disponibilidad de la historia clínica establecidos en la "
    "normatividad colombiana (Ministerio de Salud y Protección Social, "
    "1999).")
justified(doc,
    "En el caso particular de la Clínica Junical, el personal de "
    "ventanilla debe diligenciar de forma repetitiva los datos de cada "
    "solicitud, lo que genera demoras en la atención y dificulta el "
    "seguimiento del estado de cada caso. La ausencia de un registro "
    "centralizado y auditable impide identificar en qué etapa se "
    "encuentra una solicitud y verificar el cumplimiento de los "
    "tiempos de respuesta institucionales (Pressman y Maxim, 2020).")

heading2(doc, "1.1. Formulación del problema")
justified(doc,
    "¿De qué manera el diseño, desarrollo e implementación de un "
    "sistema de información web puede optimizar el registro, control y "
    "seguimiento de las solicitudes de historias clínicas en la Clínica "
    "Junical, garantizando la integridad de los datos y agilizando la "
    "atención en ventanilla?")
page_break(doc)

# ── 2. JUSTIFICACIÓN ───────────────────────────────────────────
heading1(doc, "2. Justificación")
justified(doc,
    "El presente proyecto se justifica por la necesidad real de "
    "modernizar uno de los procesos administrativos más sensibles de "
    "la Clínica Junical: la gestión de las solicitudes de historias "
    "clínicas. Contar con una herramienta tecnológica dedicada permite "
    "eliminar las deficiencias propias del manejo manual y garantizar "
    "un registro confiable, organizado y auditable, en concordancia "
    "con los principios establecidos por la normatividad colombiana "
    "(Ministerio de Salud y Protección Social, 2017).")
justified(doc,
    "Los beneficios esperados se reflejan en la reducción de los "
    "tiempos de atención al usuario, la disminución de errores humanos "
    "en la captura de datos, la posibilidad de consultar el estado de "
    "cada solicitud en tiempo real y la generación automática de "
    "reportes que apoyen la toma de decisiones administrativas "
    "(Sommerville, 2021). Los beneficiarios directos son los usuarios "
    "que acuden a la clínica para solicitar sus historias clínicas, el "
    "personal de ventanilla que dispondrá de una herramienta para "
    "agilizar su labor y la administración de la institución, que "
    "contará con información trazable para el control de gestión.")
justified(doc,
    "Desde el ámbito académico, el proyecto permite aplicar de manera "
    "integral los conocimientos adquiridos durante la formación "
    "profesional en áreas como análisis de requerimientos, diseño de "
    "software, bases de datos, desarrollo web y seguridad de la "
    "información (Pressman y Maxim, 2020). Adicionalmente, fortalece "
    "el vínculo entre la Universidad Piloto de Colombia y las "
    "instituciones del sector salud de la región del Alto Magdalena.")
justified(doc,
    "El proyecto contribuye al cumplimiento del Objetivo de Desarrollo "
    "Sostenible número tres, denominado Salud y Bienestar, en su meta "
    "relacionada con el fortalecimiento de la capacidad de las "
    "instituciones de salud para proveer servicios de calidad y con "
    "información confiable a la población (Organización de las "
    "Naciones Unidas, 2015).")
page_break(doc)

# ── 3. OBJETIVOS ───────────────────────────────────────────────
heading1(doc, "3. Objetivos")

heading2(doc, "3.1. General")
justified(doc,
    "Diseñar, desarrollar e implementar un sistema de información web "
    "que optimice el registro, control y seguimiento de historias "
    "clínicas, garantizando la integridad de los datos y agilizando la "
    "atención en ventanilla.")

heading2(doc, "3.2. Específicos")
justified(doc,
    "Analizar los procesos actuales de registro y gestión de historias "
    "clínicas, con el fin de identificar falencias y necesidades que "
    "permitan optimizar la atención en ventanilla.")
justified(doc,
    "Diseñar la arquitectura y estructura del sistema web, asegurando "
    "un flujo eficiente para el registro, control y seguimiento de la "
    "información clínica.")
justified(doc,
    "Desarrollar los módulos principales del sistema, orientados a "
    "mejorar la gestión de historias clínicas y la trazabilidad de los "
    "pacientes.")
justified(doc,
    "Implementar el sistema web como herramienta de apoyo en la "
    "gestión de historias clínicas, facilitando el acceso rápido y "
    "organizado a la información durante la atención al usuario.")
page_break(doc)

# ── 4. MARCO REFERENCIAL ───────────────────────────────────────
heading1(doc, "4. Marco referencial")

heading2(doc, "4.1. Antecedentes")
justified(doc,
    "En el ámbito internacional, diversas investigaciones han "
    "evidenciado que la implementación de sistemas de información "
    "hospitalaria mejora la oportunidad de la atención, la integridad "
    "de los registros y la capacidad de respuesta de las instituciones "
    "frente a las exigencias regulatorias (Laudon y Laudon, 2020). En "
    "particular, los sistemas orientados a la gestión administrativa "
    "del expediente clínico han demostrado disminuir los errores "
    "humanos asociados al manejo manual y reducir los tiempos de "
    "consulta (Pressman y Maxim, 2020).")
justified(doc,
    "A nivel nacional, varias entidades del sector salud han migrado "
    "sus procesos administrativos desde herramientas ofimáticas hacia "
    "soluciones web que permiten la consulta concurrente, la auditoría "
    "y la generación automatizada de reportes, atendiendo las "
    "disposiciones de la Resolución 839 de 2017 (Ministerio de Salud "
    "y Protección Social, 2017). Estas experiencias previas sirven "
    "como referente metodológico para el presente proyecto.")

heading2(doc, "4.2. Marco teórico")
justified(doc,
    "El proyecto se sustenta en la teoría de los sistemas de "
    "información, entendidos como el conjunto de componentes "
    "interrelacionados que recopilan, procesan, almacenan y distribuyen "
    "información para apoyar la toma de decisiones, la coordinación y "
    "el control en una organización (Laudon y Laudon, 2020).")
justified(doc,
    "Asimismo, se apoya en los principios de la ingeniería del "
    "software, que plantea el desarrollo de aplicaciones bajo un ciclo "
    "de vida que comprende las fases de análisis de requerimientos, "
    "diseño, implementación, pruebas y mantenimiento (Sommerville, "
    "2021). En este caso, se adopta un modelo iterativo e incremental "
    "que permite refinar los entregables a partir de la "
    "retroalimentación del usuario final (Pressman y Maxim, 2020).")
justified(doc,
    "En materia de arquitectura, el sistema se concibe bajo el modelo "
    "cliente-servidor, en el cual un componente cliente, representado "
    "por una aplicación web ejecutada en el navegador, solicita "
    "servicios a un componente servidor que procesa la información y "
    "gestiona la persistencia de los datos en una base de datos "
    "relacional (Sommerville, 2021).")
justified(doc,
    "Finalmente, los conceptos de autenticación, autorización basada "
    "en roles y registro de auditoría se constituyen como mecanismos "
    "fundamentales para garantizar la confidencialidad, integridad y "
    "trazabilidad de la información en los sistemas de salud (Laudon "
    "y Laudon, 2020).")

heading2(doc, "4.3. Marco conceptual")
justified(doc,
    "Historia clínica: documento privado, obligatorio y sometido a "
    "reserva, en el que se registran cronológicamente las condiciones "
    "de salud del paciente, los actos médicos y los demás "
    "procedimientos ejecutados por el equipo de salud que interviene "
    "en su atención (Ministerio de Salud y Protección Social, 1999).")
justified(doc,
    "Trazabilidad: capacidad de seguir y documentar la historia, "
    "ubicación y aplicación de un elemento mediante registros "
    "identificados a lo largo del tiempo (Sommerville, 2021).")
justified(doc,
    "Sistema de información web: aplicación que se ejecuta en un "
    "servidor y a la cual se accede desde un navegador a través de "
    "internet o de una red interna, ofreciendo servicios a múltiples "
    "usuarios de manera concurrente (Laudon y Laudon, 2020).")
justified(doc,
    "Bitácora o registro de auditoría: registro cronológico de las "
    "acciones ejecutadas dentro de un sistema, en el que se identifica "
    "el usuario responsable, la fecha, la hora y el detalle de la "
    "operación, con fines de control y trazabilidad (Pressman y Maxim, "
    "2020).")
justified(doc,
    "Usuario de ventanilla: funcionario encargado de atender las "
    "solicitudes presentadas por los usuarios de la institución, "
    "responsable de registrar, dar seguimiento y entregar los "
    "documentos correspondientes (Ministerio de Salud y Protección "
    "Social, 2017).")

heading2(doc, "4.4. Marco legal")
justified(doc,
    "El proyecto se desarrolla dentro del marco normativo colombiano "
    "vigente en materia de gestión de historias clínicas y protección "
    "de datos personales. En orden jerárquico se destacan las "
    "siguientes disposiciones:")
justified(doc,
    "Constitución Política de Colombia de 1991, artículo 15, que "
    "consagra el derecho a la intimidad personal y al habeas data "
    "(Asamblea Nacional Constituyente, 1991).")
justified(doc,
    "Ley Estatutaria 1751 de 2015, por medio de la cual se regula el "
    "derecho fundamental a la salud y se dictan otras disposiciones "
    "(Congreso de la República de Colombia, 2015).")
justified(doc,
    "Ley 1581 de 2012, por la cual se dictan disposiciones generales "
    "para la protección de datos personales (Congreso de la República "
    "de Colombia, 2012).")
justified(doc,
    "Decreto 1377 de 2013, que reglamenta parcialmente la Ley 1581 de "
    "2012 (Presidencia de la República de Colombia, 2013).")
justified(doc,
    "Resolución 1995 de 1999 del Ministerio de Salud, por la cual se "
    "establecen normas para el manejo de la historia clínica "
    "(Ministerio de Salud y Protección Social, 1999).")
justified(doc,
    "Resolución 839 de 2017 del Ministerio de Salud y Protección "
    "Social, que modifica la Resolución 1995 de 1999 en lo relativo al "
    "manejo, custodia, tiempo de retención y disposición final de la "
    "historia clínica (Ministerio de Salud y Protección Social, 2017).")

heading2(doc, "4.5. Marco geográfico")
justified(doc,
    "El proyecto se desarrolla en la Clínica Junical, institución "
    "prestadora de servicios de salud ubicada en la ciudad de "
    "Girardot, departamento de Cundinamarca, Colombia. Girardot se "
    "encuentra en la región del Alto Magdalena, a aproximadamente "
    "ciento treinta y cuatro kilómetros al suroeste de Bogotá, y "
    "constituye uno de los principales centros prestadores de "
    "servicios de salud para la región (Alcaldía Municipal de "
    "Girardot, 2023).")

heading2(doc, "4.6. Marco demográfico")
justified(doc,
    "Los usuarios directos del sistema son los funcionarios del área "
    "administrativa de la Clínica Junical, particularmente el personal "
    "de ventanilla encargado de la recepción y entrega de historias "
    "clínicas, así como el personal con perfil de supervisión. Los "
    "usuarios indirectos son los pacientes y demás personas que "
    "solicitan acceso a las historias clínicas, comprendiendo una "
    "población de todas las edades, géneros y condiciones "
    "socioeconómicas, propias del régimen contributivo y subsidiado "
    "del Sistema General de Seguridad Social en Salud (Hernández et "
    "al., 2014).")
page_break(doc)

# ── 5. DELIMITACIÓN ────────────────────────────────────────────
heading1(doc, "5. Delimitación")
justified(doc,
    "Delimitación temática: el alcance del proyecto se concentra en "
    "el proceso administrativo de solicitud, control y entrega de "
    "historias clínicas en ventanilla. No se aborda el contenido "
    "asistencial de las historias ni el módulo de atención médica "
    "propiamente dicho (Sommerville, 2021).")
justified(doc,
    "Delimitación espacial: el sistema se implementa en la sede "
    "principal de la Clínica Junical, ubicada en la ciudad de "
    "Girardot, departamento de Cundinamarca.")
justified(doc,
    "Delimitación temporal: el desarrollo del proyecto se realiza "
    "durante el periodo correspondiente a la práctica profesional del "
    "estudiante, con una duración aproximada de seis meses, "
    "comprendidos en el año dos mil veintiséis.")
justified(doc,
    "Delimitación funcional: el sistema contempla los módulos de "
    "gestión de pacientes, solicitudes, usuarios, bitácora de "
    "auditoría y reportes. No se incluyen integraciones con sistemas "
    "de facturación, historia clínica electrónica asistencial ni "
    "gestión documental externa (Pressman y Maxim, 2020).")
page_break(doc)

# ── 6. DISEÑO METODOLÓGICO ─────────────────────────────────────
heading1(doc, "6. Diseño metodológico")
justified(doc,
    "El presente proyecto se desarrolla bajo un enfoque de "
    "investigación aplicada, de tipo descriptivo y propositivo, en "
    "el cual se parte del estudio de una situación real para diseñar "
    "e implementar una solución concreta a una necesidad identificada "
    "(Hernández et al., 2014).")
justified(doc,
    "La metodología corresponde a un ciclo de vida de desarrollo de "
    "software de carácter iterativo e incremental, que permite "
    "ajustar los entregables a partir de la retroalimentación continua "
    "del usuario final (Sommerville, 2021). Se compone de las "
    "siguientes fases:")
justified(doc,
    "Fase de análisis: comprende el levantamiento de requerimientos "
    "mediante observación directa, entrevistas con el personal de "
    "ventanilla y revisión de los formatos utilizados en el proceso "
    "actual. Como resultado se obtiene el documento de requerimientos "
    "funcionales y no funcionales (Pressman y Maxim, 2020).")
justified(doc,
    "Fase de diseño: a partir de los requerimientos identificados, se "
    "elabora la arquitectura del sistema, el modelo de datos y los "
    "prototipos de las interfaces de usuario (Sommerville, 2021).")
justified(doc,
    "Fase de desarrollo: se realiza la implementación de los módulos "
    "del sistema de manera incremental, integrando progresivamente "
    "cada componente al producto final.")
justified(doc,
    "Fase de pruebas: se ejecutan pruebas funcionales, de integración "
    "y de aceptación con la participación del personal de ventanilla, "
    "verificando el cumplimiento de los requerimientos (Pressman y "
    "Maxim, 2020).")
justified(doc,
    "Fase de implementación: comprende la puesta en marcha del "
    "sistema en el entorno real de la clínica, la capacitación del "
    "personal usuario y la entrega de la documentación técnica y de "
    "usuario.")
justified(doc,
    "Las técnicas de recolección de información empleadas son la "
    "observación directa, la entrevista no estructurada con el "
    "personal involucrado y la revisión documental de los formatos "
    "y registros utilizados actualmente (Hernández et al., 2014). El "
    "tipo de investigación es cualitativo aplicado, con énfasis en el "
    "desarrollo de un producto tecnológico funcional.")
page_break(doc)

# ── 7. RECURSOS ────────────────────────────────────────────────
heading1(doc, "7. Recursos")

heading2(doc, "7.1. Humano")
justified(doc,
    "El recurso humano del proyecto está conformado por el estudiante "
    "practicante, Miguel Ángel Piraquive Pachón, responsable del "
    "análisis, diseño, desarrollo, pruebas e implementación del "
    "sistema. El proyecto cuenta además con el acompañamiento del "
    "docente tutor de práctica profesional de la Universidad Piloto "
    "de Colombia, encargado de la revisión metodológica y académica, "
    "y del tutor empresarial designado por la Clínica Junical, quien "
    "orienta al practicante en el contexto institucional. El personal "
    "de ventanilla de la clínica participa como usuario final y "
    "aporta información clave durante las fases de análisis y "
    "validación (Pressman y Maxim, 2020).")

heading2(doc, "7.2. Material")
justified(doc,
    "Los recursos materiales empleados comprenden un equipo de "
    "cómputo personal del estudiante con las características técnicas "
    "requeridas para el desarrollo de software, así como las "
    "herramientas de programación necesarias: editor de código, "
    "sistema de control de versiones, gestor de bases de datos "
    "relacional y entorno de ejecución para aplicaciones web "
    "(Sommerville, 2021). Adicionalmente, se requiere acceso a la "
    "infraestructura tecnológica de la Clínica Junical para la fase "
    "de implementación, así como conexión a internet y material de "
    "oficina para la elaboración de la documentación.")
page_break(doc)

# ── 8. DESARROLLO PROPUESTA ────────────────────────────────────
heading1(doc, "8. Desarrollo propuesta")
justified(doc,
    "La propuesta de solución consiste en el diseño, desarrollo e "
    "implementación de un sistema de información web denominado "
    "Sistema de Gestión y Trazabilidad de Historias Clínicas "
    "(SIGET-HC), orientado a la administración integral del proceso "
    "de solicitud, control y entrega de historias clínicas en la "
    "Clínica Junical (Laudon y Laudon, 2020).")
justified(doc,
    "El sistema se concibe bajo una arquitectura cliente-servidor, "
    "con una capa de presentación accesible desde un navegador web y "
    "una capa de servicios responsable del procesamiento de la "
    "información y de la persistencia de los datos en una base de "
    "datos relacional (Sommerville, 2021). La autenticación de los "
    "usuarios se realiza mediante un mecanismo seguro basado en "
    "credenciales con almacenamiento cifrado, y el acceso a las "
    "funcionalidades se controla mediante roles diferenciados para "
    "administrador y personal de ventanilla (Pressman y Maxim, 2020).")
justified(doc,
    "El sistema está conformado por los siguientes módulos principales:")
justified(doc,
    "Módulo de gestión de pacientes: permite registrar, consultar y "
    "actualizar la información básica de los pacientes que solicitan "
    "su historia clínica, evitando la duplicidad de registros y "
    "garantizando la integridad de los datos.")
justified(doc,
    "Módulo de solicitudes: gestiona el ciclo de vida completo de "
    "cada solicitud, desde su creación hasta su entrega, registrando "
    "los estados intermedios y permitiendo consultar en cualquier "
    "momento la situación actual de cada caso.")
justified(doc,
    "Módulo de usuarios: administra las cuentas de los funcionarios "
    "que acceden al sistema, controlando los perfiles, los permisos y "
    "la seguridad de las credenciales.")
justified(doc,
    "Módulo de bitácora: registra de manera automática y cronológica "
    "todas las acciones realizadas dentro del sistema, identificando "
    "el usuario responsable, la fecha, la hora y el detalle de la "
    "operación, dando cumplimiento a los principios de trazabilidad "
    "(Ministerio de Salud y Protección Social, 2017).")
justified(doc,
    "Módulo de reportes: genera reportes operativos y estadísticos "
    "que apoyan la toma de decisiones administrativas, permitiendo la "
    "exportación de la información en formatos estándar.")
justified(doc,
    "Para la implementación del sistema se contempla la realización "
    "de pruebas funcionales con datos reales, así como la "
    "capacitación del personal de ventanilla, asegurando la "
    "apropiación adecuada de la herramienta. Finalmente, se entrega "
    "la documentación técnica y el manual de usuario correspondiente "
    "(Sommerville, 2021).")
page_break(doc)

# ── 9. CONCLUSIONES ────────────────────────────────────────────
heading1(doc, "9. Conclusiones")
justified(doc,
    "El desarrollo del presente proyecto permitió analizar de manera "
    "detallada el proceso administrativo de gestión de historias "
    "clínicas en la Clínica Junical, identificando los puntos "
    "críticos y las oportunidades de mejora a partir de la "
    "observación directa y la interacción con el personal de "
    "ventanilla.")
justified(doc,
    "El diseño y desarrollo del sistema se ajustaron a los "
    "requerimientos identificados, ofreciendo una solución "
    "organizada, segura y trazable que reemplaza el manejo manual "
    "mediante archivos planos y disminuye los riesgos de "
    "inconsistencia y pérdida de información, en concordancia con "
    "los principios de la ingeniería del software (Pressman y Maxim, "
    "2020).")
justified(doc,
    "La implementación del sistema en el entorno real de la clínica "
    "evidencia una mejora en los tiempos de atención y en la calidad "
    "de los registros, así como una mayor confianza del personal "
    "usuario en la herramienta, lo que confirma la pertinencia de la "
    "solución propuesta.")
justified(doc,
    "Desde el punto de vista académico, el proyecto permitió "
    "consolidar los conocimientos adquiridos durante la formación "
    "profesional, articulando de manera práctica los conceptos de "
    "ingeniería de software, bases de datos, desarrollo web y "
    "seguridad de la información en una solución concreta y útil "
    "para una institución del sector salud (Sommerville, 2021).")
page_break(doc)

# ── 10. RECOMENDACIONES ────────────────────────────────────────
heading1(doc, "10. Recomendaciones")
justified(doc,
    "Se recomienda a la Clínica Junical mantener un programa "
    "permanente de capacitación al personal de ventanilla sobre el "
    "uso del sistema y sobre las buenas prácticas de manejo y "
    "custodia de la información clínica, en cumplimiento de la "
    "normatividad vigente (Ministerio de Salud y Protección Social, "
    "2017).")
justified(doc,
    "Se sugiere realizar copias de seguridad periódicas de la base "
    "de datos del sistema, almacenándolas en un sitio seguro y "
    "diferente al servidor principal, con el fin de garantizar la "
    "continuidad del servicio frente a eventos imprevistos "
    "(Sommerville, 2021).")
justified(doc,
    "Para futuras fases del proyecto se recomienda evaluar la "
    "integración del sistema con módulos asistenciales y con el "
    "sistema de facturación de la institución, lo que permitiría "
    "ampliar el alcance de la solución y aprovechar al máximo la "
    "información registrada (Laudon y Laudon, 2020).")
justified(doc,
    "Se recomienda asimismo establecer un proceso formal de "
    "actualización y mantenimiento del sistema, definiendo "
    "responsables y procedimientos para la atención de "
    "requerimientos de cambio y la incorporación de nuevas "
    "funcionalidades (Pressman y Maxim, 2020).")
justified(doc,
    "Finalmente, se sugiere a la Universidad Piloto de Colombia "
    "Seccional del Alto Magdalena fortalecer los convenios con "
    "instituciones del sector salud de la región, dado que "
    "constituyen un escenario propicio para que los estudiantes "
    "desarrollen proyectos con impacto social real y consoliden "
    "competencias profesionales útiles para el contexto local.")
page_break(doc)

# ── REFERENCIAS ────────────────────────────────────────────────
heading1(doc, "Referencias")
referencias = [
    "Alcaldía Municipal de Girardot. (2023). Información general del "
    "municipio de Girardot, Cundinamarca. "
    "https://www.girardot-cundinamarca.gov.co",

    "Asamblea Nacional Constituyente. (1991). Constitución Política "
    "de Colombia. Gaceta Constitucional No. 116.",

    "Congreso de la República de Colombia. (2012). Ley 1581 de 2012. "
    "Por la cual se dictan disposiciones generales para la "
    "protección de datos personales. Diario Oficial No. 48.587.",

    "Congreso de la República de Colombia. (2015). Ley Estatutaria "
    "1751 de 2015. Por medio de la cual se regula el derecho "
    "fundamental a la salud y se dictan otras disposiciones. Diario "
    "Oficial No. 49.427.",

    "Hernández, R., Fernández, C., y Baptista, P. (2014). "
    "Metodología de la investigación (6.ª ed.). McGraw-Hill.",

    "Instituto Colombiano de Normas Técnicas y Certificación. "
    "(2022). Compendio de normas técnicas colombianas sobre "
    "documentación. ICONTEC.",

    "Laudon, K. C., y Laudon, J. P. (2020). Sistemas de información "
    "gerencial (16.ª ed.). Pearson Educación.",

    "Ministerio de Salud y Protección Social de Colombia. (1999). "
    "Resolución 1995 de 1999. Por la cual se establecen normas para "
    "el manejo de la historia clínica.",

    "Ministerio de Salud y Protección Social de Colombia. (2017). "
    "Resolución 839 de 2017. Por la cual se modifica la Resolución "
    "1995 de 1999 y se dictan otras disposiciones.",

    "Organización de las Naciones Unidas. (2015). Objetivos de "
    "Desarrollo Sostenible. Agenda 2030 para el Desarrollo "
    "Sostenible. https://www.un.org/sustainabledevelopment/es/",

    "Presidencia de la República de Colombia. (2013). Decreto 1377 "
    "de 2013. Por el cual se reglamenta parcialmente la Ley 1581 de "
    "2012.",

    "Pressman, R. S., y Maxim, B. R. (2020). Ingeniería del "
    "software: un enfoque práctico (9.ª ed.). McGraw-Hill.",

    "Sommerville, I. (2021). Ingeniería de software (10.ª ed.). "
    "Pearson Educación.",
]
for r in referencias:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(12)
    pf.left_indent = Cm(1.25)
    pf.first_line_indent = Cm(-1.25)
    run = p.add_run(r)
    run.font.name = FONT
    run.font.size = Pt(12)

page_break(doc)

# ── ANEXOS ─────────────────────────────────────────────────────
heading1(doc, "Anexos")
justified(doc,
    "Anexo A. Manual de usuario del sistema SIGET-HC.")
justified(doc,
    "Anexo B. Matriz de requerimientos funcionales y no funcionales "
    "del proyecto.")
justified(doc,
    "Anexo C. Documento de arquitectura del sistema.")
justified(doc,
    "Anexo D. Evidencias de pruebas funcionales y de aceptación "
    "realizadas con el personal de ventanilla.")

output = (r"c:\Users\migue\Desktop\proyecto de practicas"
          r"\PROYECTO_GRADO_SIGET-HC_v2.docx")
doc.save(output)
print(f"Documento generado: {output}")
