# -*- coding: utf-8 -*-
"""
Genera el cronograma de actividades del proyecto SIGET-HC en un documento
independiente. La tabla puede copiarse y pegarse directamente en el
documento principal del proyecto.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Arial"


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def style_cell(cell, text, *, bold=False, size=10, color="000000",
               bg=None, align=WD_ALIGN_PARAGRAPH.LEFT,
               valign=WD_ALIGN_VERTICAL.CENTER):
    cell.vertical_alignment = valign
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    # limpiar runs previos
    for r in list(p.runs):
        r.text = ""
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    if bg:
        set_cell_bg(cell, bg)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[idx])


# ═══════════════════════════════════════════════════════════════
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)
    # orientación apaisada para que quepa la tabla
    from docx.enum.section import WD_ORIENT
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width, s.page_height = s.page_height, s.page_width

# Título de sección
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("13.1.4. Cronograma de actividades")
run.font.name = FONT
run.font.size = Pt(14)
run.bold = True

# Párrafo descriptivo
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(12)
run = p.add_run(
    "El siguiente cronograma de actividades detalla las tareas ejecutadas "
    "para el desarrollo del proyecto Sistema de Gestión y Trazabilidad de "
    "Historias Clínicas (SIGET-HC), incluyendo la fase, la actividad "
    "específica, la duración asignada en meses, las semanas estimadas y "
    "el responsable de su realización.")
run.font.name = FONT
run.font.size = Pt(12)

# ── Tabla 1: Cronograma detallado por fases ───────────────────
headers = ["Fase", "Actividad", "Duración (meses)",
           "Semanas", "Responsable"]
data = [
    ("Análisis",
     "Levantamiento de requerimientos en ventanilla",
     "Mes 1", "Semanas 1–2", "Estudiante practicante"),
    ("Análisis",
     "Entrevistas con el personal de ventanilla",
     "Mes 1", "Semanas 2–3", "Estudiante practicante"),
    ("Análisis",
     "Análisis del proceso actual y documentación de "
     "requerimientos funcionales y no funcionales",
     "Mes 1", "Semanas 3–4", "Estudiante practicante / Tutor empresarial"),
    ("Diseño",
     "Diseño de la arquitectura del sistema (cliente-servidor)",
     "Mes 2", "Semanas 5–6", "Estudiante practicante"),
    ("Diseño",
     "Modelado de la base de datos y diccionario de datos",
     "Mes 2", "Semanas 6–7", "Estudiante practicante"),
    ("Diseño",
     "Prototipado de interfaces de usuario y validación con "
     "ventanilla",
     "Mes 2", "Semanas 7–8", "Estudiante practicante / Personal de "
     "ventanilla"),
    ("Desarrollo",
     "Configuración del entorno y proyecto base (Django + React)",
     "Mes 3", "Semana 9", "Estudiante practicante"),
    ("Desarrollo",
     "Desarrollo del módulo de usuarios y autenticación JWT",
     "Mes 3", "Semanas 9–10", "Estudiante practicante"),
    ("Desarrollo",
     "Desarrollo del módulo de pacientes e integración con "
     "Verifik",
     "Mes 3", "Semanas 10–12", "Estudiante practicante"),
    ("Desarrollo",
     "Desarrollo del módulo de solicitudes y flujo de estados",
     "Mes 4", "Semanas 13–15", "Estudiante practicante"),
    ("Desarrollo",
     "Desarrollo del módulo de bitácora de auditoría",
     "Mes 4", "Semana 16", "Estudiante practicante"),
    ("Desarrollo",
     "Desarrollo del módulo de reportes y exportación a Excel",
     "Mes 5", "Semanas 17–18", "Estudiante practicante"),
    ("Pruebas",
     "Pruebas funcionales internas de cada módulo",
     "Mes 5", "Semanas 18–19", "Estudiante practicante"),
    ("Pruebas",
     "Pruebas de integración entre módulos",
     "Mes 5", "Semana 20", "Estudiante practicante"),
    ("Pruebas",
     "Pruebas de aceptación con el personal de ventanilla",
     "Mes 6", "Semana 21", "Estudiante practicante / Personal de "
     "ventanilla"),
    ("Implementación",
     "Despliegue del sistema en la infraestructura de la "
     "Clínica Junical",
     "Mes 6", "Semana 22", "Estudiante practicante / Tutor empresarial"),
    ("Implementación",
     "Capacitación al personal de ventanilla",
     "Mes 6", "Semana 23", "Estudiante practicante"),
    ("Cierre",
     "Entrega de la documentación técnica, manual de usuario "
     "e informe final",
     "Mes 6", "Semana 24", "Estudiante practicante / Tutor académico"),
]

table = doc.add_table(rows=1 + len(data), cols=len(headers))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

NAVY = "173A6B"
SKY = "CFEAF7"
ZEBRA = "F4F8FC"

# Encabezado
for idx, h in enumerate(headers):
    style_cell(table.rows[0].cells[idx], h, bold=True, size=11,
               color="FFFFFF", bg=NAVY,
               align=WD_ALIGN_PARAGRAPH.CENTER)

# Filas (cebra + cabecera de fase coloreada)
prev_fase = None
for i, row_data in enumerate(data, start=1):
    fase = row_data[0]
    is_new_fase = fase != prev_fase
    bg_row = SKY if is_new_fase else (ZEBRA if i % 2 == 0 else None)
    for j, value in enumerate(row_data):
        bold = (j == 0 and is_new_fase)
        align = (WD_ALIGN_PARAGRAPH.LEFT if j == 1
                 else WD_ALIGN_PARAGRAPH.CENTER)
        style_cell(table.rows[i].cells[j], value, bold=bold,
                   size=10, bg=bg_row, align=align)
    prev_fase = fase

# Anchos de columna (suma ≈ 25.7 cm en orientación apaisada)
set_col_widths(table, [2.8, 8.5, 3.0, 3.2, 5.5])

# Bordes
tbl = table._tbl
tblPr = tbl.tblPr
tblBorders = OxmlElement("w:tblBorders")
for border_name in ("top", "left", "bottom", "right",
                    "insideH", "insideV"):
    b = OxmlElement(f"w:{border_name}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "4")
    b.set(qn("w:color"), "9CB3D1")
    tblBorders.append(b)
tblPr.append(tblBorders)

# Pie de tabla
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Tabla 1. Cronograma de actividades del proyecto SIGET-HC.")
run.font.name = FONT
run.font.size = Pt(10)
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Fuente: elaboración propia.")
run.font.name = FONT
run.font.size = Pt(10)
run.italic = True

output = (r"c:\Users\migue\Desktop\proyecto de practicas"
          r"\CRONOGRAMA_SIGET-HC.docx")
doc.save(output)
print(f"Cronograma generado: {output}")
